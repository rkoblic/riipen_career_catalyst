#!/usr/bin/env python3
"""
Convert Markdown files to Riipen-branded .docx documents.

Uses Riipen brand guidelines:
- Typography: Crimson Pro (headings/title), DM Sans (body/subheadings)
- Colors: Dark Blue (#050C2B), Orange (#FF7C0A), Electric Blue (#2454FF)
- Layout: Riipen logo top-right, address header, consistent spacing

Usage:
    python3 scripts/md_to_docx.py input.md [output.docx] [--logo path/to/logo.png]
"""

import sys
import os
import re
import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Riipen Brand Colors ──────────────────────────────────────────────
DARK_BLUE = RGBColor(0x05, 0x0C, 0x2A)
ORANGE = RGBColor(0xFF, 0x7C, 0x0A)
ELECTRIC_BLUE = RGBColor(0x24, 0x54, 0xFF)
GREEN = RGBColor(0x18, 0x73, 0x3E)
GREY = RGBColor(0xAC, 0xAC, 0xAC)
LIGHT_GREY = RGBColor(0xFA, 0xFA, 0xFA)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)

# Table header background
TABLE_HEADER_BG = "FF7C0A"
TABLE_ALT_ROW_BG = "F5F5F5"
TABLE_BORDER_COLOR = "ACACAC"

# ── Font Settings ────────────────────────────────────────────────────
# All document text uses DM Sans (matching the Riipen document template)
TITLE_FONT = "DM Sans"
HEADING_FONT = "DM Sans"
BODY_FONT = "DM Sans"
SUBHEADING_FONT = "DM Sans"

TITLE_SIZE = Pt(28)
H1_SIZE = Pt(17)
H2_SIZE = Pt(14)
H3_SIZE = Pt(12)
H4_SIZE = Pt(11)
BODY_SIZE = Pt(10)
SMALL_SIZE = Pt(9)

# ── Address ──────────────────────────────────────────────────────────
RIIPEN_ADDRESS = "1200-555 West Hastings Street Vancouver, BC, V6B 4N6"


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table, color="000000", size="4"):
    """Set borders on a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    """Set spacing on a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:before="{before}" w:after="{after}"/>')
    if line:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    # Remove existing spacing
    for existing in pPr.findall(qn('w:spacing')):
        pPr.remove(existing)
    pPr.append(spacing)


def format_run(run, font_name=BODY_FONT, size=BODY_SIZE, color=DARK_BLUE,
               bold=False, italic=False, underline=False):
    """Apply formatting to a run."""
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    run.underline = underline
    # Ensure font applies to complex script text too
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def add_formatted_text(paragraph, text, font_name=BODY_FONT, size=BODY_SIZE,
                       color=DARK_BLUE, bold=False, italic=False, underline=False):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    format_run(run, font_name, size, color, bold, italic, underline)
    return run


def parse_inline_markdown(paragraph, text, font_name=BODY_FONT, size=BODY_SIZE,
                          color=DARK_BLUE, base_bold=False):
    """Parse inline markdown (bold, italic, links, code) and add runs to paragraph."""
    # Pattern for: **bold**, *italic*, `code`, [link](url), ***bold italic***
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'   # ***bold italic***
        r'|(\*\*(.+?)\*\*)'      # **bold**
        r'|(\*(.+?)\*)'          # *italic*
        r'|(`(.+?)`)'            # `code`
        r'|(\[(.+?)\]\((.+?)\))' # [link](url)
    )

    last_end = 0
    for match in pattern.finditer(text):
        # Add text before the match
        if match.start() > last_end:
            add_formatted_text(paragraph, text[last_end:match.start()],
                             font_name, size, color, bold=base_bold)

        if match.group(2):  # ***bold italic***
            add_formatted_text(paragraph, match.group(2),
                             font_name, size, color, bold=True, italic=True)
        elif match.group(4):  # **bold**
            add_formatted_text(paragraph, match.group(4),
                             font_name, size, color, bold=True)
        elif match.group(6):  # *italic*
            add_formatted_text(paragraph, match.group(6),
                             font_name, size, color, italic=True)
        elif match.group(8):  # `code`
            add_formatted_text(paragraph, match.group(8),
                             "Courier New", Pt(9.5), DARK_BLUE)
        elif match.group(10):  # [link](url)
            link_text = match.group(10)
            # link_url = match.group(11)  # We can't easily add hyperlinks with python-docx
            add_formatted_text(paragraph, link_text,
                             font_name, size, ORANGE, underline=True)

        last_end = match.end()

    # Add remaining text
    if last_end < len(text):
        add_formatted_text(paragraph, text[last_end:],
                         font_name, size, color, bold=base_bold)


def add_header(doc, logo_path=None):
    """Add Riipen header with logo and address."""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # If we have a logo, add it right-aligned
    if logo_path and os.path.exists(logo_path):
        logo_para = header.paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = logo_para.add_run()
        run.add_picture(logo_path, width=Inches(1.2))

        # Address below logo
        addr_para = header.add_paragraph()
        addr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_formatted_text(addr_para, RIIPEN_ADDRESS,
                         BODY_FONT, SMALL_SIZE, GREY)
    else:
        # Just address, right-aligned
        addr_para = header.paragraphs[0]
        addr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_formatted_text(addr_para, RIIPEN_ADDRESS,
                         BODY_FONT, SMALL_SIZE, GREY)

    set_paragraph_spacing(addr_para, before=0, after=120)


def add_title(doc, title_text):
    """Add the document title in Riipen style — orange accent."""
    para = doc.add_paragraph()
    add_formatted_text(para, title_text, HEADING_FONT, H1_SIZE, ORANGE, bold=True)
    set_paragraph_spacing(para, before=360, after=0)
    return para


def add_heading(doc, text, level=1):
    """Add a heading in Riipen style."""
    para = doc.add_paragraph()

    if level == 1:
        add_formatted_text(para, text, HEADING_FONT, H1_SIZE, DARK_BLUE, bold=True)
        set_paragraph_spacing(para, before=360, after=120)
    elif level == 2:
        add_formatted_text(para, text, SUBHEADING_FONT, H2_SIZE, DARK_BLUE, bold=True)
        set_paragraph_spacing(para, before=280, after=100)
    elif level == 3:
        add_formatted_text(para, text, SUBHEADING_FONT, H3_SIZE, DARK_BLUE, bold=True)
        set_paragraph_spacing(para, before=240, after=80)
    else:
        add_formatted_text(para, text, SUBHEADING_FONT, H4_SIZE, DARK_BLUE, bold=True)
        set_paragraph_spacing(para, before=200, after=60)

    return para


def add_body_paragraph(doc, text):
    """Add a body text paragraph with inline markdown parsing."""
    para = doc.add_paragraph()
    parse_inline_markdown(para, text)
    set_paragraph_spacing(para, before=0, after=120, line=276)
    return para


def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    para = doc.add_paragraph()
    # Set indent for nesting (values in twips: 1 inch = 1440 twips)
    indent_left = 720 + level * 360
    indent_hanging = 360

    pPr = para._p.get_or_add_pPr()
    ind = parse_xml(
        f'<w:ind {nsdecls("w")} w:left="{indent_left}" w:hanging="{indent_hanging}"/>'
    )
    for existing in pPr.findall(qn('w:ind')):
        pPr.remove(existing)
    pPr.append(ind)

    # Add bullet character
    bullet_char = "\u2022" if level == 0 else "\u25E6"
    add_formatted_text(para, f"{bullet_char}  ", BODY_FONT, BODY_SIZE, ORANGE, bold=True)
    parse_inline_markdown(para, text)
    set_paragraph_spacing(para, before=0, after=60, line=276)
    return para


def add_numbered_item(doc, text, number):
    """Add a numbered list item."""
    para = doc.add_paragraph()
    # Values in twips: 1 inch = 1440 twips
    indent_left = 720
    indent_hanging = 360

    pPr = para._p.get_or_add_pPr()
    ind = parse_xml(
        f'<w:ind {nsdecls("w")} w:left="{indent_left}" w:hanging="{indent_hanging}"/>'
    )
    for existing in pPr.findall(qn('w:ind')):
        pPr.remove(existing)
    pPr.append(ind)

    add_formatted_text(para, f"{number}.  ", BODY_FONT, BODY_SIZE, DARK_BLUE, bold=True)
    parse_inline_markdown(para, text)
    set_paragraph_spacing(para, before=0, after=60, line=276)
    return para


def calc_column_widths(headers, rows, total_width):
    """Calculate proportional column widths based on content length."""
    num_cols = len(headers)

    # Find the max content length per column (header + all rows)
    max_lengths = []
    for col_idx in range(num_cols):
        col_max = len(headers[col_idx])
        for row in rows:
            if col_idx < len(row):
                col_max = max(col_max, len(row[col_idx]))
        max_lengths.append(col_max)

    # Apply a minimum width (so tiny columns aren't unreadably narrow)
    MIN_CHARS = 4
    max_lengths = [max(ml, MIN_CHARS) for ml in max_lengths]

    # Use sqrt scaling so long-text columns don't dominate too aggressively
    import math
    scaled = [math.sqrt(ml) for ml in max_lengths]
    total_scaled = sum(scaled)

    # Convert to EMU widths proportionally
    widths = [int((s / total_scaled) * total_width) for s in scaled]

    # Correct rounding errors
    diff = total_width - sum(widths)
    widths[-1] += diff

    return widths


def add_table(doc, headers, rows):
    """Add a branded table with auto-fit column widths."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table width to full page
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="dxa" w:w="9360"/>')
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblPr.append(tblW)

    # Disable auto-layout so our widths are respected
    tblLayout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    for existing in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(existing)
    tblPr.append(tblLayout)

    set_table_borders(table, color=TABLE_BORDER_COLOR, size="4")

    # Calculate best-fit column widths (in twips; 9360 twips = 6.5in page width)
    col_widths = calc_column_widths(headers, rows, 9360)

    # Apply column widths
    for col_idx in range(len(headers)):
        for row in table.rows:
            cell = row.cells[col_idx]
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = parse_xml(
                f'<w:tcW {nsdecls("w")} w:type="dxa" w:w="{col_widths[col_idx]}"/>'
            )
            for existing in tcPr.findall(qn('w:tcW')):
                tcPr.remove(existing)
            tcPr.append(tcW)

    # Header row
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, TABLE_HEADER_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        add_formatted_text(para, header_text, SUBHEADING_FONT, Pt(10), WHITE, bold=True)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(para, before=40, after=40)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            if row_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_ROW_BG)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            parse_inline_markdown(para, cell_text, BODY_FONT, Pt(10), DARK_BLUE)
            set_paragraph_spacing(para, before=30, after=30)

    # Add spacing after table
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=120)

    return table


def add_horizontal_rule(doc):
    """Add a horizontal rule / divider."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{TABLE_BORDER_COLOR}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    set_paragraph_spacing(para, before=120, after=120)
    return para


def add_blockquote(doc, text):
    """Add a blockquote with left border styling."""
    para = doc.add_paragraph()

    # Left border and indent
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="FF7C0A"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="720"/>')
    pPr.append(ind)

    parse_inline_markdown(para, text, BODY_FONT, BODY_SIZE, GREY, base_bold=False)
    set_paragraph_spacing(para, before=120, after=120, line=276)
    return para


def add_code_block(doc, code_text):
    """Add a code block with grey background."""
    para = doc.add_paragraph()

    # Grey background via shading
    pPr = para._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    pPr.append(shading)

    # Add indent for padding effect
    ind = parse_xml(
        f'<w:ind {nsdecls("w")} w:left="432" w:right="432"/>'
    )
    pPr.append(ind)

    add_formatted_text(para, code_text, "Courier New", Pt(9), DARK_BLUE)
    set_paragraph_spacing(para, before=60, after=60)
    return para


def parse_markdown_table(lines):
    """Parse a markdown table into headers and rows."""
    headers = []
    rows = []

    for i, line in enumerate(lines):
        line = line.strip()
        if not line.startswith('|'):
            continue

        cells = [c.strip() for c in line.split('|')[1:-1]]

        # Skip separator row (e.g., |---|---|)
        if all(re.match(r'^[-:]+$', c) for c in cells):
            continue

        if not headers:
            headers = cells
        else:
            # Pad row to match header count
            while len(cells) < len(headers):
                cells.append('')
            rows.append(cells[:len(headers)])

    return headers, rows


def parse_markdown(md_text):
    """Parse markdown text into a list of block elements."""
    lines = md_text.split('\n')
    blocks = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Empty line
        if not stripped:
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            blocks.append(('heading', level, text))
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[-*_]{3,}\s*$', stripped):
            blocks.append(('hr',))
            i += 1
            continue

        # Table
        if '|' in stripped and i + 1 < len(lines) and '|' in lines[i + 1]:
            table_lines = []
            while i < len(lines) and '|' in lines[i].strip():
                table_lines.append(lines[i])
                i += 1
            headers, rows = parse_markdown_table(table_lines)
            if headers:
                blocks.append(('table', headers, rows))
            continue

        # Code block
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            blocks.append(('code', '\n'.join(code_lines)))
            continue

        # Blockquote
        if stripped.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(re.sub(r'^>\s?', '', lines[i].strip()))
                i += 1
            blocks.append(('blockquote', ' '.join(quote_lines)))
            continue

        # Unordered list
        list_match = re.match(r'^(\s*)([-*+])\s+(.+)$', stripped)
        if list_match:
            items = []
            while i < len(lines):
                item_match = re.match(r'^(\s*)([-*+])\s+(.+)$', lines[i])
                if item_match:
                    indent = len(item_match.group(1))
                    level = indent // 2
                    items.append((level, item_match.group(3).strip()))
                    i += 1
                elif lines[i].strip() == '':
                    i += 1
                    # Check if list continues
                    if i < len(lines) and re.match(r'^(\s*)([-*+])\s+', lines[i]):
                        continue
                    break
                else:
                    break
            blocks.append(('bullet_list', items))
            continue

        # Ordered list
        num_match = re.match(r'^(\s*)\d+[.)]\s+(.+)$', stripped)
        if num_match:
            items = []
            while i < len(lines):
                item_match = re.match(r'^(\s*)\d+[.)]\s+(.+)$', lines[i])
                if item_match:
                    items.append(item_match.group(2).strip())
                    i += 1
                elif lines[i].strip() == '':
                    i += 1
                    if i < len(lines) and re.match(r'^\s*\d+[.)]\s+', lines[i]):
                        continue
                    break
                else:
                    break
            blocks.append(('numbered_list', items))
            continue

        # Regular paragraph (may span multiple lines)
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_stripped = lines[i].strip()
            if (not next_stripped or
                re.match(r'^#{1,6}\s+', next_stripped) or
                re.match(r'^[-*_]{3,}\s*$', next_stripped) or
                re.match(r'^(\s*)([-*+])\s+', next_stripped) or
                re.match(r'^\d+[.)]\s+', next_stripped) or
                next_stripped.startswith('```') or
                next_stripped.startswith('>') or
                ('|' in next_stripped and i + 1 < len(lines) and '|' in lines[i + 1].strip())):
                break
            para_lines.append(next_stripped)
            i += 1
        blocks.append(('paragraph', ' '.join(para_lines)))

    return blocks


def convert_md_to_docx(md_path, output_path, logo_path=None):
    """Convert a markdown file to a Riipen-branded docx."""
    md_text = Path(md_path).read_text(encoding='utf-8')
    blocks = parse_markdown(md_text)

    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # ── Set default font ──
    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    style.font.color.rgb = DARK_BLUE
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), BODY_FONT)
    rFonts.set(qn('w:hAnsi'), BODY_FONT)
    rFonts.set(qn('w:cs'), BODY_FONT)

    # ── Header ──
    add_header(doc, logo_path)

    # ── Process blocks ──
    title_added = False
    hr_count = 0

    for block in blocks:
        btype = block[0]

        if btype == 'heading':
            level = block[1]
            text = block[2]
            # First H1 becomes the document title
            if level == 1 and not title_added:
                add_title(doc, text)
                title_added = True
            else:
                add_heading(doc, text, level)

        elif btype == 'paragraph':
            text = block[1]
            add_body_paragraph(doc, text)

        elif btype == 'bullet_list':
            for level, text in block[1]:
                add_bullet(doc, text, level)

        elif btype == 'numbered_list':
            for idx, text in enumerate(block[1], 1):
                add_numbered_item(doc, text, idx)

        elif btype == 'table':
            headers = block[1]
            rows = block[2]
            add_table(doc, headers, rows)

        elif btype == 'code':
            for code_line in block[1].split('\n'):
                add_code_block(doc, code_line)

        elif btype == 'blockquote':
            add_blockquote(doc, block[1])

        elif btype == 'hr':
            hr_count += 1
            if hr_count <= 1:
                add_horizontal_rule(doc)

    # ── Footer with decorative element ──
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(footer_para, "Riipen", HEADING_FONT, Pt(8), ORANGE, bold=True)
    run = footer_para.add_run("  |  ")
    format_run(run, BODY_FONT, Pt(8), GREY)
    add_formatted_text(footer_para, "riipen.com", BODY_FONT, Pt(8), GREY)

    doc.save(output_path)
    return output_path


def find_logo(script_dir):
    """Auto-detect Riipen logo file relative to the project root."""
    # Project root is one level up from scripts/
    project_root = script_dir.parent
    candidates = [
        project_root / "Riipen_Logo.jpg",
        project_root / "Riipen_Logo.png",
        project_root / "riipen_logo.jpg",
        project_root / "riipen_logo.png",
        project_root / "logo.png",
        project_root / "logo.jpg",
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def main():
    parser = argparse.ArgumentParser(
        description='Convert Markdown to Riipen-branded .docx'
    )
    parser.add_argument('input', help='Input markdown file path')
    parser.add_argument('output', nargs='?', help='Output docx file path (default: same name as input with .docx extension)')
    parser.add_argument('--logo', help='Path to Riipen logo image file (auto-detected if not specified)', default=None)
    parser.add_argument('--no-logo', action='store_true', help='Omit logo from header')

    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: Input file '{input_path}' not found.")
        sys.exit(1)

    if args.output:
        output_path = args.output
    else:
        output_path = str(input_path.with_suffix('.docx'))

    logo_path = None
    if not args.no_logo:
        logo_path = args.logo or find_logo(Path(__file__).resolve().parent)
        if logo_path:
            print(f"Using logo: {logo_path}")

    result = convert_md_to_docx(str(input_path), output_path, logo_path)
    print(f"Created: {result}")


if __name__ == '__main__':
    main()
