#!/usr/bin/env python3
"""
Combine all markdown content files for a week into a single .docx for review.

Not branded — just clean formatting with proper headers, body text, and page breaks
between pages. Intended for pasting into Google Docs for stakeholder review.

Usage:
    python3 scripts/combine_week_content.py <week_number> [output.docx]

Examples:
    python3 scripts/combine_week_content.py 1
    python3 scripts/combine_week_content.py 01 weekly_content/week01/week01-review.docx
"""

import sys
import os
import re
import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml

# ── Constants ──

FONT_NAME = "DM Sans"
DARK_BLUE = RGBColor(0x05, 0x0C, 0x2A)
MUTED = RGBColor(0x6B, 0x72, 0x80)


def parse_args():
    parser = argparse.ArgumentParser(description="Combine week content into a single .docx")
    parser.add_argument("week", help="Week number (e.g., 1 or 01)")
    parser.add_argument("output", nargs="?", default=None, help="Output .docx path")
    return parser.parse_args()


def strip_frontmatter(text):
    """Remove YAML frontmatter from markdown."""
    if text.startswith("---"):
        end = text.find("---", 3)
        if end != -1:
            return text[end + 3:].strip()
    return text.strip()


def extract_frontmatter(text):
    """Extract YAML frontmatter fields as a dict."""
    meta = {}
    if text.startswith("---"):
        end = text.find("---", 3)
        if end != -1:
            block = text[3:end].strip()
            for line in block.split("\n"):
                if ":" in line:
                    key, _, val = line.partition(":")
                    meta[key.strip()] = val.strip()
    return meta


def set_run_font(run, size=11, bold=False, italic=False, color=DARK_BLUE):
    """Apply font formatting to a run."""
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Ensure DM Sans is used
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {{{qn("w")}: "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}} w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:ascii"), FONT_NAME)
        rFonts.set(qn("w:hAnsi"), FONT_NAME)


def add_formatted_text(paragraph, text, size=11, bold=False, italic=False, color=DARK_BLUE):
    """Add a run with formatting to a paragraph."""
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return run


def process_inline(paragraph, text, size=11, color=DARK_BLUE):
    """Process markdown inline formatting (bold, italic, bold+italic, links)."""
    # Pattern for bold+italic, bold, italic, links, and plain text
    pattern = re.compile(
        r'\*\*\*(.+?)\*\*\*'      # ***bold italic***
        r'|\*\*(.+?)\*\*'          # **bold**
        r'|\*(.+?)\*'              # *italic*
        r'|\[([^\]]+)\]\(([^)]+)\)' # [text](url)
        r'|`([^`]+)`'              # `code`
        r'|([^*\[`]+(?:[*\[`](?![*\[`])[^*\[`]*)*)'  # plain text
    )

    for m in pattern.finditer(text):
        if m.group(1):  # bold italic
            add_formatted_text(paragraph, m.group(1), size=size, bold=True, italic=True, color=color)
        elif m.group(2):  # bold
            add_formatted_text(paragraph, m.group(2), size=size, bold=True, color=color)
        elif m.group(3):  # italic
            add_formatted_text(paragraph, m.group(3), size=size, italic=True, color=color)
        elif m.group(4):  # link
            add_formatted_text(paragraph, m.group(4), size=size, color=color)
        elif m.group(6):  # code
            add_formatted_text(paragraph, m.group(6), size=size, color=color)
        elif m.group(7):  # plain
            add_formatted_text(paragraph, m.group(7), size=size, color=color)


VIDEO_SCRIPT_COLOR = RGBColor(0x8B, 0x5C, 0x0A)  # dark amber for video script label


def add_video_script_inline(doc, script_content):
    """Render a video script inline where the [VIDEO:] placeholder was."""
    # Label
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    add_formatted_text(p, "▶ VIDEO SCRIPT", size=10, bold=True, color=VIDEO_SCRIPT_COLOR)

    # Render script body as indented muted paragraphs (skip H1 title)
    for line in script_content.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            # Use the H1 as a subtitle for the script
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            add_formatted_text(p, stripped[2:].strip(), size=11, bold=True, italic=True, color=MUTED)
            continue
        if stripped in ("---", "***", "___"):
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        process_inline(p, stripped, size=11, color=MUTED)

    # End label
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    add_formatted_text(p, "▶ END VIDEO SCRIPT", size=10, bold=True, color=VIDEO_SCRIPT_COLOR)


def add_page_content(doc, content, page_num, title, video_scripts=None):
    """Add a single page's markdown content to the document.

    video_scripts: dict mapping page number (str) to list of stripped video
    script contents, to be inlined at [VIDEO:] placeholders.
    """
    video_scripts = video_scripts or {}
    page_videos = list(video_scripts.get(str(page_num), []))  # copy so we can pop
    lines = content.split("\n")
    i = 0
    in_list = False
    list_ordered = False
    list_counter = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Skip horizontal rules
        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # H1 — page title (Heading 1)
        if stripped.startswith("# ") and not stripped.startswith("## "):
            heading_text = stripped[2:].strip()
            p = doc.add_heading(heading_text, level=1)
            for run in p.runs:
                set_run_font(run, size=18, bold=True, color=DARK_BLUE)
            i += 1
            continue

        # H2 — section header (Heading 2)
        if stripped.startswith("## "):
            heading_text = stripped[3:].strip()
            p = doc.add_heading(heading_text, level=2)
            for run in p.runs:
                set_run_font(run, size=14, bold=True, color=DARK_BLUE)
            i += 1
            continue

        # H3 — subsection (Heading 3)
        if stripped.startswith("### "):
            heading_text = stripped[4:].strip()
            p = doc.add_heading(heading_text, level=3)
            for run in p.runs:
                set_run_font(run, size=12, bold=True, color=DARK_BLUE)
            i += 1
            continue

        # Blockquote — render as indented italic
        if stripped.startswith("> "):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_line = lines[i].strip()
                if quote_line == ">":
                    quote_lines.append("")
                else:
                    quote_lines.append(quote_line[2:] if quote_line.startswith("> ") else quote_line[1:])
                i += 1
            quote_text = "\n".join(quote_lines)
            for para_text in quote_text.split("\n"):
                if para_text.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    process_inline(p, para_text.strip(), size=11, color=MUTED)
                else:
                    # Empty line in blockquote
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_after = Pt(2)
            continue

        # Unordered list
        if stripped.startswith("- ") or stripped.startswith("* "):
            item_text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            process_inline(p, item_text, size=11)
            i += 1
            continue

        # Ordered list
        ol_match = re.match(r"^(\d+)\.\s+(.+)", stripped)
        if ol_match:
            item_text = ol_match.group(2).strip()
            p = doc.add_paragraph(style="List Number")
            process_inline(p, item_text, size=11)
            i += 1
            continue

        # Checkbox list items
        if stripped.startswith("- [ ] ") or stripped.startswith("- [x] "):
            item_text = stripped[6:].strip()
            p = doc.add_paragraph(style="List Bullet")
            process_inline(p, item_text, size=11)
            i += 1
            continue

        # Markdown table — detect by pipe row followed by separator row
        if stripped.startswith("|") and stripped.endswith("|") and i + 1 < len(lines):
            next_stripped = lines[i + 1].strip()
            is_separator = (
                next_stripped.startswith("|")
                and next_stripped.endswith("|")
                and set(next_stripped.replace("|", "").replace(":", "").replace(" ", "")) <= {"-"}
                and "-" in next_stripped
            )
            if is_separator:
                # Collect all table rows
                header_cells = [c.strip() for c in stripped.strip("|").split("|")]
                i += 2  # skip header and separator
                body_rows = []
                while i < len(lines):
                    row = lines[i].strip()
                    if not (row.startswith("|") and row.endswith("|")):
                        break
                    body_rows.append([c.strip() for c in row.strip("|").split("|")])
                    i += 1

                num_cols = len(header_cells)
                table = doc.add_table(rows=1 + len(body_rows), cols=num_cols)
                table.style = "Light Grid Accent 1"
                table.autofit = True

                # Header row
                for col_idx, cell_text in enumerate(header_cells):
                    cell = table.rows[0].cells[col_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    process_inline(p, cell_text, size=11)
                    for run in p.runs:
                        run.font.bold = True

                # Body rows
                for row_idx, row in enumerate(body_rows, start=1):
                    for col_idx in range(num_cols):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = ""
                        cell_text = row[col_idx] if col_idx < len(row) else ""
                        p = cell.paragraphs[0]
                        process_inline(p, cell_text, size=11)

                # Spacer paragraph after table
                doc.add_paragraph()
                continue

        # VIDEO placeholder — inline the video script if available
        if stripped.startswith("[VIDEO:"):
            p = doc.add_paragraph()
            add_formatted_text(p, stripped, size=10, italic=True, color=MUTED)
            if page_videos:
                add_video_script_inline(doc, page_videos.pop(0))
            i += 1
            continue

        # Other placeholders — render as grey italic
        if stripped.startswith("[INTERACTIVE:") or \
           stripped.startswith("[CURATED LINK:") or stripped.startswith("[TEMPLATE:") or \
           stripped.startswith("[LINKED RESOURCE:"):
            p = doc.add_paragraph()
            add_formatted_text(p, stripped, size=10, italic=True, color=MUTED)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        process_inline(p, stripped, size=11)
        i += 1


def inline_video_scripts_in_md(content, videos):
    """Return the page content with each [VIDEO:] placeholder followed by the
    corresponding video script rendered as a markdown blockquote, so the
    script appears inline in the combined markdown output."""
    videos = list(videos)  # copy so we can pop
    lines = content.split("\n")
    output_lines = []
    for line in lines:
        output_lines.append(line)
        stripped = line.strip()
        if stripped.startswith("[VIDEO:") and videos:
            script = videos.pop(0)
            output_lines.append("")
            output_lines.append("> **▶ Video script**")
            output_lines.append(">")
            for script_line in script.split("\n"):
                s = script_line.strip()
                if s in ("---", "***", "___"):
                    continue
                if s.startswith("# "):
                    output_lines.append(f"> *{s[2:].strip()}*")
                    output_lines.append(">")
                elif s:
                    output_lines.append(f"> {s}")
                else:
                    output_lines.append(">")
            output_lines.append("")
    return "\n".join(output_lines)


def build_combined_markdown(pages_data, video_scripts, week_num, summary_parts):
    """Build a single combined markdown string from all pages, with video
    scripts inlined at their [VIDEO:] placeholders."""
    parts = []
    parts.append(f"# Week {int(week_num)} — Content Review\n")
    parts.append(f"\n**RIIPEN CAREER CATALYST** — {' + '.join(summary_parts)}\n")
    parts.append("\n---\n")

    for idx, (page_num, _title, content, _filepath) in enumerate(pages_data):
        parts.append(f"\n**PAGE {page_num}**\n\n")
        page_videos = video_scripts.get(str(page_num), [])
        parts.append(inline_video_scripts_in_md(content, page_videos))
        if idx < len(pages_data) - 1:
            parts.append("\n\n---\n")

    return "".join(parts)


def main():
    args = parse_args()

    # Normalize week number
    week_num = args.week.zfill(2)
    week_dir = Path(f"weekly_content/week{week_num}")

    if not week_dir.exists():
        print(f"Error: {week_dir} does not exist.")
        sys.exit(1)

    # Separate content pages from video scripts and interactives
    def page_sort_key(p):
        m = re.match(r"page(\d+)", p.name)
        return (int(m.group(1)) if m else 9999, p.name)
    all_md = sorted(week_dir.glob("page*.md"), key=page_sort_key)
    content_files = [f for f in all_md if "-video.md" not in f.name and "-interactive.md" not in f.name]
    video_files = [f for f in all_md if "-video.md" in f.name]

    if not content_files:
        print(f"Error: No content files found in {week_dir}/")
        sys.exit(1)

    # Build video script lookup: page number -> list of script contents
    video_scripts = {}
    for vf in video_files:
        raw = vf.read_text(encoding="utf-8")
        meta = extract_frontmatter(raw)
        page_num = meta.get("page", "0")
        script_content = strip_frontmatter(raw)
        video_scripts.setdefault(page_num, []).append(script_content)
    video_count = len(video_files)

    # Determine output path
    if args.output:
        output_path = Path(args.output)
    else:
        output_path = week_dir / f"week{week_num}-content-review.docx"

    # Create document
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style.font.color.rgb = DARK_BLUE

    # Set heading styles
    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = FONT_NAME
        heading_style.font.color.rgb = DARK_BLUE

    # Compute summary info (used by both markdown and docx outputs)
    page_count = len(content_files)
    summary_parts = [f"{page_count} pages"]
    if video_count:
        summary_parts.append(f"{video_count} video scripts")

    # Collect page data once, used by both outputs
    pages_data = []
    for idx, filepath in enumerate(content_files):
        raw = filepath.read_text(encoding="utf-8")
        meta = extract_frontmatter(raw)
        content = strip_frontmatter(raw)
        page_num = meta.get("page", str(idx + 1))
        title = meta.get("title", filepath.stem)
        pages_data.append((page_num, title, content, filepath))

    # Build and save combined markdown alongside the docx
    combined_md = build_combined_markdown(pages_data, video_scripts, week_num, summary_parts)
    md_output_path = output_path.with_suffix(".md")
    md_output_path.write_text(combined_md, encoding="utf-8")

    # Add cover info to the docx
    p = doc.add_paragraph()
    add_formatted_text(p, f"Week {int(week_num)} — Content Review", size=24, bold=True)
    p = doc.add_paragraph()
    add_formatted_text(p, "RIIPEN CAREER CATALYST", size=11, bold=True, color=MUTED)
    p = doc.add_paragraph()
    add_formatted_text(p, " + ".join(summary_parts), size=11, color=MUTED)
    doc.add_page_break()

    # Process each content page (video scripts are inlined at their placeholders)
    for idx, (page_num, title, content, _filepath) in enumerate(pages_data):
        # Page label
        p = doc.add_paragraph()
        add_formatted_text(p, f"PAGE {page_num}", size=10, bold=True, color=MUTED)

        # Add content, passing video scripts for inline insertion
        add_page_content(doc, content, page_num, title, video_scripts=video_scripts)

        # Page break between pages (not after the last one)
        if idx < len(pages_data) - 1:
            doc.add_page_break()

    # Save docx
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    print(f"Also created: {md_output_path}")
    print(f"  Content pages: {page_count}")
    print(f"  Video scripts: {video_count} (inlined at placeholders)")
    for f in content_files:
        print(f"    - {f.name}")
    for vf in video_files:
        print(f"    - {vf.name} (inlined)")


if __name__ == "__main__":
    main()
