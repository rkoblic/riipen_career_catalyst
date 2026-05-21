#!/usr/bin/env python3
"""
Build the annotated "Formatting professional deliverables" sample DOCX
referenced from weekly_content/week06/page5-formatting-professional-deliverables.md.

The sample is a fictional student team's short final-deliverable-style report
written for GreenPath Solutions (the recurring fictional firm from weeks 1-2).
Inline annotations (orange italic) point out heading hierarchy, white space,
table format, chart formatting, callout box use, and terminology consistency.

Output: weekly_content/week06/page5-formatting-example.docx
Also writes a one-off chart PNG into weekly_content/week06/ for embedding.

Run: python3 scripts/build_formatting_example.py
"""

from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
WEEK_DIR = ROOT / "weekly_content" / "week06"
OUT_DOCX = WEEK_DIR / "page5-formatting-example.docx"
CHART_PNG = WEEK_DIR / "page5-formatting-example-chart.png"

# Professional report palette (NOT Riipen brand — this is a student team's doc)
INK = RGBColor(0x1F, 0x2D, 0x3D)      # body text
MUTED = RGBColor(0x5A, 0x6A, 0x7A)    # secondary text
ACCENT = RGBColor(0x1E, 0x5F, 0x8F)   # report accent (blue)
HEADER_BG = "1E5F8F"                  # table header shading
CALLOUT_BG = "EAF3FA"                 # soft blue callout fill
CALLOUT_BORDER = "1E5F8F"

# Annotation styling — visually obvious "this is meta-content"
ANNOT = RGBColor(0xC2, 0x57, 0x0C)    # burnt orange
ANNOT_PREFIX = "ANNOTATION  ·  "

REPORT_FONT = "Calibri"  # widely available; student-deliverable feel


# ── low-level helpers ──────────────────────────────────────────────────

def _set_run_font(run, name, size, color, bold=False, italic=False):
    run.font.name = name
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)


def _spacing(paragraph, before=0, after=120, line=276):
    pPr = paragraph._p.get_or_add_pPr()
    for existing in pPr.findall(qn('w:spacing')):
        pPr.remove(existing)
    spacing = parse_xml(
        f'<w:spacing {nsdecls("w")} w:before="{before}" w:after="{after}" '
        f'w:line="{line}" w:lineRule="auto"/>'
    )
    pPr.append(spacing)


def _cell_shading(cell, color_hex):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _table_borders(table, color="C5CCD3", size="4"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


# ── content helpers ────────────────────────────────────────────────────

def add_title(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    _set_run_font(r, REPORT_FONT, Pt(22), ACCENT, bold=True)
    _spacing(p, before=0, after=80, line=240)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run_font(r, REPORT_FONT, Pt(12), MUTED, italic=True)
    _spacing(p, before=0, after=320, line=276)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    r = p.add_run(text)
    _set_run_font(r, REPORT_FONT, Pt(14), ACCENT, bold=True)
    _spacing(p, before=280, after=100, line=276)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    r = p.add_run(text)
    _set_run_font(r, REPORT_FONT, Pt(11.5), INK, bold=True)
    _spacing(p, before=200, after=60, line=276)
    return p


def add_body(doc, text, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run_font(r, REPORT_FONT, Pt(11), INK, italic=italic)
    _spacing(p, before=0, after=140, line=288)
    return p


def add_caption(doc, text):
    """Centered, small, muted caption used for table titles and figure titles."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_run_font(r, REPORT_FONT, Pt(10), MUTED, bold=True)
    _spacing(p, before=120, after=60, line=276)
    return p


def add_annotation(doc, label, text):
    """Inline annotation — orange italic, small, with a leading marker."""
    p = doc.add_paragraph()
    # Left border to make it visually distinct as a side-note
    pPr = p._p.get_or_add_pPr()
    border = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="18" w:space="6" w:color="C2570C"/>'
        f'</w:pBdr>'
    )
    pPr.append(border)
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="180"/>')
    pPr.append(ind)
    r1 = p.add_run(ANNOT_PREFIX + label.upper())
    _set_run_font(r1, REPORT_FONT, Pt(9), ANNOT, bold=True)
    r2 = p.add_run("  " + text)
    _set_run_font(r2, REPORT_FONT, Pt(9), ANNOT, italic=True)
    _spacing(p, before=40, after=200, line=240)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    _set_run_font(r, REPORT_FONT, Pt(11), INK)
    _spacing(p, before=0, after=80, line=276)
    return p


def add_segment_table(doc):
    """Donor segment table — title above (added separately), header row shaded."""
    rows = [
        ["Donor segment", "Average gift", "First-year retention", "Multi-year retention"],
        ["First-time donors", "$85", "28%", "14%"],
        ["Recurring donors", "$145", "71%", "58%"],
        ["Major donors ($500+)", "$1,250", "84%", "76%"],
    ]
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _table_borders(table, color="C5CCD3", size="4")

    # widths roughly proportional to content
    widths = [Inches(1.9), Inches(1.1), Inches(1.5), Inches(1.5)]
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.width = widths[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            r = p.add_run(val)
            if i == 0:
                _set_run_font(r, REPORT_FONT, Pt(10), RGBColor(0xFF, 0xFF, 0xFF), bold=True)
                _cell_shading(cell, HEADER_BG)
            else:
                _set_run_font(r, REPORT_FONT, Pt(10.5), INK)
            _spacing(p, before=60, after=60, line=264)
    return table


def add_callout_box(doc, heading, body):
    """A 1x1 shaded table acting as a callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    _cell_shading(cell, CALLOUT_BG)

    # subtle full border in accent color
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="{CALLOUT_BORDER}"/>'
        f'  <w:left w:val="single" w:sz="18" w:space="0" w:color="{CALLOUT_BORDER}"/>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="{CALLOUT_BORDER}"/>'
        f'  <w:right w:val="single" w:sz="8" w:space="0" w:color="{CALLOUT_BORDER}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

    # First paragraph already exists in cell
    p1 = cell.paragraphs[0]
    r = p1.add_run(heading)
    _set_run_font(r, REPORT_FONT, Pt(10.5), ACCENT, bold=True)
    _spacing(p1, before=100, after=60, line=264)

    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    _set_run_font(r2, REPORT_FONT, Pt(10.5), INK)
    _spacing(p2, before=0, after=100, line=276)
    return table


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="6" w:color="D5DCE3"/>'
        f'</w:pBdr>'
    )
    pPr.append(pbdr)
    _spacing(p, before=120, after=120)
    return p


def add_chart_image(doc):
    """Generate a clean bar chart and insert as an image."""
    fig, ax = plt.subplots(figsize=(6.2, 3.3))
    segments = ["First-time", "Recurring", "Major"]
    revenue = [24, 51, 25]  # share of total revenue, %
    bars = ax.bar(segments, revenue, color=["#9CB7CC", "#1E5F8F", "#3A87B9"], width=0.55)
    for b, v in zip(bars, revenue):
        ax.text(b.get_x() + b.get_width() / 2, v + 1.2, f"{v}%",
                ha="center", va="bottom", fontsize=10, color="#1F2D3D")
    ax.set_ylabel("Share of annual donor revenue (%)", fontsize=10, color="#1F2D3D")
    ax.set_xlabel("Donor segment", fontsize=10, color="#1F2D3D")
    ax.set_ylim(0, 65)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.tick_params(colors="#5A6A7A")
    ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, _: f"{int(x)}%"))
    plt.tight_layout()
    fig.savefig(CHART_PNG, dpi=180, bbox_inches="tight")
    plt.close(fig)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(CHART_PNG), width=Inches(5.8))
    _spacing(p, before=80, after=20)


# ── document assembly ──────────────────────────────────────────────────

def build():
    doc = Document()

    # Page setup — 1" margins all round
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Default font for the document
    style = doc.styles["Normal"]
    style.font.name = REPORT_FONT
    style.font.size = Pt(11)

    # ─── Reader note (cover) ───────────────────────────────────────
    cover = doc.add_paragraph()
    cover_r = cover.add_run(
        "HOW TO READ THIS EXAMPLE  ·  This is an annotated sample of a "
        "short final-deliverable-style report. The orange notes in the "
        "left margin point out specific formatting choices that match the "
        "standards on the course page. The report has been shortened — a "
        "real deliverable would be longer — but the standards apply at "
        "any length. The fictional firm in this example, GreenPath "
        "Solutions, is the same firm referenced earlier in the course."
    )
    _set_run_font(cover_r, REPORT_FONT, Pt(9.5), MUTED, italic=True)
    _spacing(cover, before=0, after=320, line=264)

    # ─── Title ─────────────────────────────────────────────────────
    add_title(doc, "Donor Engagement Strategy")
    add_subtitle(
        doc,
        "A recommendations report for GreenPath Solutions  ·  "
        "Prepared by the River Valley Consulting Team  ·  May 2026",
    )
    add_annotation(
        doc, "Heading hierarchy",
        "One H1 used only for the document title. Major sections below "
        "use H2; subsections use H3. The hierarchy is consistent throughout."
    )

    # ─── Executive Summary ─────────────────────────────────────────
    add_h2(doc, "Executive Summary")
    add_body(
        doc,
        "GreenPath Solutions' individual donor base grew 12 percent in "
        "2025, but recurring giving lagged peer firms in the sustainability "
        "sector. This report identifies three priorities for strengthening "
        "donor retention and recommends a phased approach the team can "
        "implement within the next two quarters."
    )
    add_body(
        doc,
        "Our analysis draws on twelve months of donor data, four "
        "stakeholder interviews, and a benchmark scan of comparable "
        "sustainability consultancies. The findings consistently point to "
        "one strategic shift: GreenPath's donor relationships are healthy "
        "at the entry stage but thin out before reaching recurring status."
    )
    add_body(
        doc,
        "The three priorities described below focus on what GreenPath "
        "can do in the conversion window between a donor's first gift and "
        "their second."
    )
    add_annotation(
        doc, "White space and paragraph length",
        "Paragraphs run four to six lines and are separated by consistent "
        "spacing. Walls of text are split at natural breaks rather than "
        "left as a single block."
    )

    # ─── Background ────────────────────────────────────────────────
    add_h2(doc, "Background")
    add_body(
        doc,
        "GreenPath Solutions, a Portland-based sustainability consulting "
        "firm, supports a network of individual and corporate donors who "
        "fund its community programming arm. Individual donors currently "
        "contribute roughly 38 percent of program revenue."
    )
    add_body(
        doc,
        "Over the past year, GreenPath's leadership has identified donor "
        "retention as a strategic gap. While first-time donor acquisition "
        "is healthy, the firm sees significantly fewer of these donors "
        "return for a second gift than peer firms report."
    )

    # ─── Methodology ───────────────────────────────────────────────
    add_h2(doc, "Methodology")
    add_body(
        doc,
        "Our team conducted three streams of research between March and "
        "April 2026:"
    )
    add_bullet(doc, "Quantitative review of twelve months of GreenPath donor records, including gift size, frequency, and source channel.")
    add_bullet(doc, "Four stakeholder interviews with GreenPath staff: the Director of Development, two program leads, and the volunteer coordinator.")
    add_bullet(doc, "A sector benchmark comparing donor engagement practices at five comparable sustainability consultancies, drawing on publicly available materials.")

    # ─── Findings ──────────────────────────────────────────────────
    add_h2(doc, "Findings")

    add_h3(doc, "Donor segments")
    add_body(
        doc,
        "GreenPath's donor base divides into three distinct segments with "
        "different giving patterns and retention rates. Table 1 summarizes "
        "the segments and the patterns associated with each."
    )
    add_caption(doc, "Table 1.  Donor segments at GreenPath Solutions, 2025")
    add_segment_table(doc)
    add_body(
        doc,
        "The retention gap between first-time and recurring donors is the "
        "largest in the segment data. Closing even a portion of that gap "
        "represents a substantial revenue opportunity.",
        italic=False,
    )
    add_annotation(
        doc, "Table format",
        "Descriptive title placed ABOVE the table (not below). Header row "
        "is visually distinct. Column widths are sized to content. The "
        "table is referenced from the surrounding text rather than left "
        "to stand on its own."
    )

    add_h3(doc, "Giving patterns")
    add_body(
        doc,
        "Recurring donors give roughly three times more annually than the "
        "average first-time donor, both because their gift sizes are "
        "higher and because their gifts repeat. Figure 1 shows the share "
        "of annual donor revenue contributed by each segment."
    )
    add_caption(doc, "Figure 1.  Share of annual donor revenue, by donor segment (2025)")
    add_chart_image(doc)
    add_body(
        doc,
        "Source: GreenPath internal donor records, March 2026."
    )
    add_annotation(
        doc, "Chart formatting",
        "Title above the chart. Axes are labeled, including units. A "
        "single accent color is used to draw the eye to the segment "
        "carrying the largest share. The data source is cited directly "
        "beneath the visual."
    )

    # ─── Key insight callout ───────────────────────────────────────
    add_callout_box(
        doc,
        "Key finding",
        "The largest revenue opportunity is not acquiring new donors. It "
        "is moving more first-time donors into the recurring segment. A "
        "five percentage point lift in first-year retention would generate "
        "more revenue than a comparable acquisition increase.",
    )
    add_annotation(
        doc, "Callout box, used sparingly",
        "A single callout box flags the report's most important finding. "
        "Used once, the callout draws attention. Used on every page, "
        "callouts lose their effect — so this report has only one."
    )

    # ─── Recommendations ───────────────────────────────────────────
    add_h2(doc, "Recommendations")
    add_body(
        doc,
        "Based on the findings above, our team recommends three "
        "priorities for the next two quarters. Each is described briefly "
        "below; full implementation detail appears in Appendix A."
    )

    add_h3(doc, "Priority 1: A structured first-90-days flow")
    add_body(
        doc,
        "Most first-time donors disengage between their first gift and "
        "the moment a typical communication cadence would re-engage them. "
        "A structured 90-day flow — thank-you call within 48 hours, "
        "impact email at day 30, invitation to a small event at day 60 — "
        "closes this gap without significant new spend."
    )

    add_h3(doc, "Priority 2: Recurring giving as the default ask")
    add_body(
        doc,
        "GreenPath's current giving page presents a one-time gift as the "
        "default, with recurring giving as a secondary choice. Reversing "
        "this default has been shown across the sector to lift recurring "
        "conversion meaningfully, with no negative effect on overall "
        "giving."
    )

    add_h3(doc, "Priority 3: Donor segmentation in the CRM")
    add_body(
        doc,
        "The current CRM does not distinguish first-time, recurring, and "
        "major donors in a way that supports differentiated communication. "
        "A modest reconfiguration would allow GreenPath's communications "
        "team to tailor messaging without adding staff capacity."
    )
    add_annotation(
        doc, "Terminology consistency",
        "Throughout the report, the same terms are used for the same "
        "things: donor segment (not group, tier, or bucket); recurring "
        "donor (not repeat, returning, or monthly); first-time donor "
        "(not new, one-time, or initial). Inconsistent terms read as "
        "careless and can confuse the reader about whether two terms "
        "mean the same thing."
    )

    # ─── Next steps ────────────────────────────────────────────────
    add_h2(doc, "Next Steps")
    add_body(
        doc,
        "We recommend GreenPath's Development team work through these "
        "priorities in sequence rather than parallel, beginning with the "
        "first-90-days flow as the highest-leverage near-term move. Our "
        "team is available to support the design of the day-1 through "
        "day-90 communication sequence in the next two weeks."
    )

    add_horizontal_rule(doc)
    closing = doc.add_paragraph()
    cr = closing.add_run(
        "River Valley Consulting Team  ·  Submitted May 2026  ·  "
        "Prepared for GreenPath Solutions"
    )
    _set_run_font(cr, REPORT_FONT, Pt(9), MUTED, italic=True)
    _spacing(closing, before=20, after=0)

    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")
    print(f"Wrote {CHART_PNG}")


if __name__ == "__main__":
    build()
