#!/usr/bin/env python3
"""
Build the one-page printable Feedback Protocol Card referenced from
weekly_content/week07/page4-receiving-feedback-well.md.

The card holds the four-step in-conversation protocol
(Listen / Clarify / Acknowledge / Decide later) with example phrases,
plus a bottom strip covering async feedback habits and what to do when
feedback isn't useful. Designed to be printed or saved by learners for
peer reviews, employer meetings, and informal feedback conversations.

Output: templates/feedback-protocol-card.docx

Run: python3 scripts/build_feedback_protocol_card.py
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT_DOCX = ROOT / "templates" / "feedback-protocol-card.docx"

# Riipen brand
FONT = "DM Sans"
ORANGE = RGBColor(0xFF, 0x7C, 0x0A)
BURNT_ORANGE = RGBColor(0xB5, 0x47, 0x08)
DARK_BLUE = RGBColor(0x05, 0x0C, 0x2A)
GREY = RGBColor(0x6B, 0x70, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

CARD_FILL = "FFF8F1"          # soft warm card background
CARD_BORDER = "FF7C0A"        # orange card border
CARD_BORDER_QUIET = "E2E4E9"  # muted card border for the bottom strip


# ── low-level helpers ──────────────────────────────────────────────────

def _set_run_font(run, name, size, color, bold=False, italic=False):
    run.font.name = name
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)


def _spacing(paragraph, before=0, after=80, line=264):
    pPr = paragraph._p.get_or_add_pPr()
    for existing in pPr.findall(qn('w:spacing')):
        pPr.remove(existing)
    spacing = parse_xml(
        f'<w:spacing {nsdecls("w")} w:before="{before}" w:after="{after}" '
        f'w:line="{line}" w:lineRule="auto"/>'
    )
    pPr.append(spacing)


def _cell_shading(cell, color_hex):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _cell_borders(cell, color_hex, sz="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def _cell_margins(cell, top=120, left=180, bottom=120, right=180):
    """Inner padding for a table cell, in twips."""
    tcPr = cell._tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="nil"/>'
        f'  <w:left w:val="nil"/>'
        f'  <w:bottom w:val="nil"/>'
        f'  <w:right w:val="nil"/>'
        f'  <w:insideH w:val="nil"/>'
        f'  <w:insideV w:val="nil"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _table_cell_spacing(table, twips=80):
    """Add spacing between cells so the cards visually separate."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    spacing = parse_xml(
        f'<w:tblCellSpacing {nsdecls("w")} w:w="{twips}" w:type="dxa"/>'
    )
    tblPr.append(spacing)


# ── content paragraph helpers ──────────────────────────────────────────

def _add_run(paragraph, text, size, color, bold=False, italic=False):
    r = paragraph.add_run(text)
    _set_run_font(r, FONT, size, color, bold=bold, italic=italic)
    return r


def _new_para(cell_or_doc, alignment=None):
    """Add a paragraph; if cell, reuse its first empty paragraph the first time."""
    if hasattr(cell_or_doc, "paragraphs") and cell_or_doc.paragraphs and not cell_or_doc.paragraphs[0].runs and getattr(cell_or_doc, "_first_used", False) is False:
        p = cell_or_doc.paragraphs[0]
        cell_or_doc._first_used = True
    else:
        p = cell_or_doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    return p


# ── card builders ──────────────────────────────────────────────────────

def add_step_card(cell, number, title, body, phrases_label=None, phrases=None):
    """Render one protocol step inside a table cell."""
    _cell_shading(cell, CARD_FILL)
    _cell_borders(cell, CARD_BORDER, sz="10")
    _cell_margins(cell, top=140, left=180, bottom=140, right=180)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Step number + title on one line
    p = cell.paragraphs[0]
    _add_run(p, f"{number}.  ", Pt(12), ORANGE, bold=True)
    _add_run(p, title.upper(), Pt(12), DARK_BLUE, bold=True)
    _spacing(p, before=0, after=40, line=252)

    # Body
    p2 = cell.add_paragraph()
    _add_run(p2, body, Pt(9), DARK_BLUE)
    _spacing(p2, before=0, after=40, line=252)

    # Phrases block (optional)
    if phrases:
        plabel = cell.add_paragraph()
        _add_run(plabel, phrases_label or "TRY:", Pt(8), BURNT_ORANGE, bold=True)
        _spacing(plabel, before=20, after=10, line=232)
        for ph in phrases:
            pph = cell.add_paragraph()
            _add_run(pph, "•  ", Pt(8.5), ORANGE, bold=True)
            _add_run(pph, ph, Pt(8.5), DARK_BLUE, italic=True)
            pPr = pph._p.get_or_add_pPr()
            ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="160" w:hanging="160"/>')
            pPr.append(ind)
            _spacing(pph, before=0, after=10, line=240)


def add_strip_card(cell, label, body_blocks, color_hex_border=CARD_BORDER_QUIET):
    """Render the bottom-strip card (async / when not useful)."""
    _cell_shading(cell, "FAFBFC")
    _cell_borders(cell, color_hex_border, sz="6")
    _cell_margins(cell, top=120, left=160, bottom=120, right=160)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    p = cell.paragraphs[0]
    _add_run(p, label.upper(), Pt(8.5), BURNT_ORANGE, bold=True)
    _spacing(p, before=0, after=40, line=232)

    for kind, text in body_blocks:
        if kind == "p":
            pp = cell.add_paragraph()
            _add_run(pp, text, Pt(8.5), DARK_BLUE)
            _spacing(pp, before=0, after=40, line=252)
        elif kind == "b":
            pp = cell.add_paragraph()
            _add_run(pp, "•  ", Pt(8.5), ORANGE, bold=True)
            _add_run(pp, text, Pt(8.5), DARK_BLUE)
            pPr = pp._p.get_or_add_pPr()
            ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="160" w:hanging="160"/>')
            pPr.append(ind)
            _spacing(pp, before=0, after=10, line=240)


# ── document assembly ──────────────────────────────────────────────────

def build():
    doc = Document()

    # Tight margins so the card fits on a single portrait page
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)

    # ─── Title ─────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    _add_run(p_title, "Feedback Protocol Card", Pt(18), ORANGE, bold=True)
    _spacing(p_title, before=0, after=20, line=232)

    p_sub = doc.add_paragraph()
    _add_run(
        p_sub,
        "A reference for receiving feedback well — in meetings, in writing, and anywhere in between.",
        Pt(9.5),
        GREY,
        italic=True,
    )
    _spacing(p_sub, before=0, after=80, line=240)

    # ─── Intro paragraph ───────────────────────────────────────────
    p_intro = doc.add_paragraph()
    _add_run(
        p_intro,
        (
            "Receiving feedback on your work is uncomfortable, and the instinct is to defend or "
            "deflect. This protocol gives you a structure that keeps you listening when the "
            "instinct is to argue — direct, indirect, written, or live."
        ),
        Pt(9.5),
        DARK_BLUE,
    )
    _spacing(p_intro, before=0, after=100, line=264)

    # ─── 2x2 grid of step cards ────────────────────────────────────
    grid = doc.add_table(rows=2, cols=2)
    grid.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid.autofit = False
    _remove_table_borders(grid)
    _table_cell_spacing(grid, twips=80)

    col_width = Inches(3.55)
    for row in grid.rows:
        for cell in row.cells:
            cell.width = col_width

    add_step_card(
        grid.cell(0, 0),
        number="1",
        title="Listen",
        body=(
            "Hear the feedback before you respond. Don't formulate your reply while the "
            "other person is still talking. Most feedback gets more specific as the person "
            "keeps explaining — the version you would have argued with after one sentence "
            "often isn't what they meant."
        ),
        phrases_label="REMIND YOURSELF:",
        phrases=[
            "Wait for the end before reacting.",
            "If you catch \"but we chose that because…\" — pause.",
            "Their first sentence is rarely their full point.",
        ],
    )

    add_step_card(
        grid.cell(0, 1),
        number="2",
        title="Clarify",
        body=(
            "Before reacting, make sure you understand what's being said. A clarifying "
            "question isn't pushing back — it shows you're taking the feedback seriously "
            "and gives you something to work with."
        ),
        phrases_label="TRY:",
        phrases=[
            "\"Can you say more about what you mean by [their term]?\"",
            "\"Is the issue with the scope, the approach, or how we've communicated it?\"",
            "\"Can you point me to a specific moment you noticed that?\"",
            "\"What would a stronger version look like to you?\"",
        ],
    )

    add_step_card(
        grid.cell(1, 0),
        number="3",
        title="Acknowledge",
        body=(
            "Name what you heard. Acknowledging isn't agreeing — you're confirming the "
            "message landed, not committing to act on it. It also gives the other person "
            "a chance to correct you."
        ),
        phrases_label="TRY:",
        phrases=[
            "\"So the concern is that [your paraphrase].\"",
            "\"What I'm hearing is [X]. Did I get that right?\"",
            "\"Is this for next week, or a longer-term consideration?\"",
            "\"That makes sense. We'll take it back and figure out what to do with it.\"",
        ],
    )

    add_step_card(
        grid.cell(1, 1),
        number="4",
        title="Decide later",
        body=(
            "Don't commit to a response in the moment. You don't have to agree with every "
            "piece of feedback, and you don't have to decide what to do with it on the spot. "
            "Your job in the conversation is to hear and understand. The decision happens "
            "after, with your team."
        ),
        phrases_label="REMIND YOURSELF:",
        phrases=[
            "Restraint now ≠ agreement later.",
            "Weigh feedback against the brief and your team's judgment.",
            "Vague feedback is harder to use — clarify first, decide second.",
        ],
    )

    # spacer before bottom strip
    spacer = doc.add_paragraph()
    _spacing(spacer, before=0, after=0, line=160)

    # ─── Bottom strip: async habits + when feedback isn't useful ───
    strip = doc.add_table(rows=1, cols=2)
    strip.alignment = WD_TABLE_ALIGNMENT.CENTER
    strip.autofit = False
    _remove_table_borders(strip)
    _table_cell_spacing(strip, twips=80)
    for cell in strip.rows[0].cells:
        cell.width = col_width

    add_strip_card(
        strip.cell(0, 0),
        label="When the feedback is written",
        body_blocks=[
            ("b", "A short comment isn't a curt one — assume the generous reading."),
            ("b", "If a comment makes you bristle, that's information about you, not the comment."),
            ("b", "You don't owe an immediate reply. Sit with it overnight."),
            ("b", "Clarify and acknowledge still apply — in writing, on your timeline."),
        ],
    )

    add_strip_card(
        strip.cell(0, 1),
        label="When feedback isn't useful",
        body_blocks=[
            (
                "p",
                "Receiving feedback well doesn't mean accepting every piece. Use CLARIFY "
                "to turn vague feedback (\"your tone needs work\") into something concrete, "
                "or to surface that there isn't anything concrete behind it.",
            ),
            (
                "p",
                "Use DECIDE LATER to weigh feedback against the brief and your team's "
                "judgment. The cost of waiting is small; the cost of arguing in the moment "
                "is the information the other person was trying to give you.",
            ),
        ],
    )

    # ─── Footer credit ─────────────────────────────────────────────
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(
        footer,
        "Riipen Career Catalyst  ·  Feedback Protocol Card",
        Pt(7.5),
        GREY,
    )
    _spacing(footer, before=100, after=0, line=200)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    build()
